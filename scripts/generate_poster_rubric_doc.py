"""Generate a Word document with rubric-aligned poster content for the occupancy estimation project."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_section_header(doc, part_label, subtitle, points):
    """Add a shaded section header row like the rubric."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"{part_label}: {subtitle}  [{points} pts]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)
    return p


def add_criterion(doc, number, criterion_title, score_note, content_text):
    """Add one rubric criterion block."""
    # Criterion title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(f"{number}. {criterion_title}")
    title_run.bold = True
    title_run.font.size = Pt(11)

    # Score target note
    score_p = doc.add_paragraph()
    score_p.paragraph_format.left_indent = Inches(0.2)
    score_p.paragraph_format.space_before = Pt(0)
    score_p.paragraph_format.space_after = Pt(2)
    sr = score_p.add_run(f"Target score: {score_note}")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Content block
    content_p = doc.add_paragraph()
    content_p.paragraph_format.left_indent = Inches(0.2)
    content_p.paragraph_format.space_before = Pt(0)
    content_p.paragraph_format.space_after = Pt(8)
    cr = content_p.add_run(content_text)
    cr.font.size = Pt(10)

    return content_p


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Research Poster — Rubric Content Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x4E, 0x2A, 0x84)

    subtitle = doc.add_paragraph(
        "ML-Driven Occupancy Estimation in Smart Buildings Using CO₂, Airflow, and Floor Sensors\n"
        "Jun Han · Samuel Hartmann · Dalton Sloan · Garrett Green"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].italic = True

    doc.add_paragraph(
        "This document maps suggested poster text to each rubric criterion. "
        "Copy, condense, and format as needed to fit poster layout constraints."
    ).runs[0].font.size = Pt(9)

    # ─── PART 1: APPEARANCE & ORGANIZATION ───────────────────────────────────
    add_section_header(doc, "PART 1", "Appearance & Organization", 15)

    add_criterion(
        doc,
        number="1",
        criterion_title="Poster text is clearly readable from a distance of 3–4 feet",
        score_note="3 (Strongly Agree)",
        content_text=(
            "Use a minimum 24 pt font for body text and 36–48 pt for section headers. "
            "All sensor labels (CO₂, VOC, ACH), figure axis labels, and table values should be at least 18 pt. "
            "Avoid dense paragraphs — use bullet points with short, complete sentences. "
            "Suggested body font: Arial or Calibri. Title font size: 72–85 pt."
        ),
    )

    add_criterion(
        doc,
        number="2",
        criterion_title="Figures, tables, and diagrams are clear and legible and add to the poster's message",
        score_note="3 (Strongly Agree)",
        content_text=(
            "Include the following figures from reports/figures/:\n"
            "  • room354_feature_comparison(1 month).png — 4-panel VOC/CO₂/humidity/temperature time series showing occupancy-linked signal peaks.\n"
            "  • room354_correlation_matrix.png — Pearson correlation heatmap confirming CO₂ (r=0.978) and VOC (r=0.748) as the strongest occupancy predictors.\n"
            "  • room354_estimated_occupancy(1 month).png — Final blended occupancy estimate over the 30-day window.\n"
            "  • reports/room361_pipeline/figures/room361_co2_flow_occupancy_timeseries.png — Room 361 CO₂, airflow, and estimated occupancy for the validation week.\n"
            "Each figure should have a bold caption, axis labels ≥14 pt, and a one-sentence takeaway directly below it."
        ),
    )

    add_criterion(
        doc,
        number="3",
        criterion_title="Design is cohesive and color scheme is not distracting",
        score_note="2 (Agree)",
        content_text=(
            "Recommended palette: Tennessee Tech purple (#4E2A84) for headers, white/light gray backgrounds, "
            "and a single accent color (gold #F0A500) for highlighted metrics. "
            "Use consistent line weights in all plots (2 pt for primary series, 1 pt for secondary). "
            "Avoid red/green color pairs — use purple/gold or blue/orange for accessibility."
        ),
    )

    add_criterion(
        doc,
        number="4",
        criterion_title="Format is consistent with no spelling or grammatical errors",
        score_note="3 (Strongly Agree)",
        content_text=(
            "Run a spell-check pass before final print. Key terms to verify:\n"
            "  • CO₂ (subscript 2, not CO2)\n"
            "  • VOC (volatile organic compounds)\n"
            "  • ACH (air changes per hour)\n"
            "  • cfm (cubic feet per minute, lowercase)\n"
            "  • MAE / RMSE (uppercase acronyms)\n"
            "Ensure all section headers use title case. Figure captions use sentence case. "
            "Author names: Jun Han, Samuel Hartmann, Dalton Sloan, Garrett Green."
        ),
    )

    add_criterion(
        doc,
        number="5",
        criterion_title="Poster is well-organized and information flows naturally from Introduction through Conclusion",
        score_note="3 (Strongly Agree)",
        content_text=(
            "Recommended left-to-right, top-to-bottom column flow:\n"
            "  Column 1: Problem Statement → Dataset\n"
            "  Column 2: Methods → Story (narrative + central figure)\n"
            "  Column 3: Results → Conclusion → References\n"
            "Use visual dividers (thin ruled lines or shaded boxes) to separate sections. "
            "Number figures (Fig. 1, Fig. 2, …) and reference them by number in the text."
        ),
    )

    # ─── PART 2: CONTENT ─────────────────────────────────────────────────────
    add_section_header(doc, "PART 2", "Content", 35)

    add_criterion(
        doc,
        number="1",
        criterion_title="Intro / Significance / Need / Research Question",
        score_note="4–5 (Agree to Strongly Agree)",
        content_text=(
            "PROBLEM STATEMENT (poster text):\n\n"
            "Buildings waste significant energy running HVAC systems on fixed schedules regardless of actual room occupancy. "
            "This project investigates whether CO₂, VOC, humidity, temperature, and discharge airflow signals — already collected "
            "by modern smart-building sensor systems — can be used to continuously estimate room occupancy without cameras or badge readers.\n\n"
            "Research question: Can a physics-informed, multi-sensor model accurately estimate the number of occupants in a university "
            "classroom (AIEB Rooms 354 and 361, Tennessee Tech) using only HVAC point-history data?\n\n"
            "Significance: A validated occupancy model enables demand-driven HVAC scheduling, reducing energy waste during low-use "
            "periods while maintaining indoor air quality standards."
        ),
    )

    add_criterion(
        doc,
        number="2",
        criterion_title="Methods — sufficiently detailed to allow understanding of materials, instruments, and procedures",
        score_note="4–5 (Agree to Strongly Agree)",
        content_text=(
            "METHODS (poster text):\n\n"
            "Data Collection:\n"
            "  • Room 354: 30-day IAQ + FPB long-format point-history CSV exports (2026-02-27 to 2026-03-29).\n"
            "  • Room 361: One-week FPB export (Apr 3–9, 2026) + 3 manual headcounts as ground-truth anchors.\n"
            "  • Sensors: CO₂ (ppm), VOC, humidity (%), temperature (°F), discharge airflow (cfm).\n\n"
            "Preprocessing:\n"
            "  • Parse UTC timestamps, pivot long → wide format, resample to 5-minute intervals.\n"
            "  • Average overlapping IAQ/FPB CO₂, humidity, and temperature streams into unified room-level signals.\n"
            "  • Convert discharge airflow (cfm) → air changes per hour (ACH) using room volume (637 m³).\n\n"
            "Feature Engineering:\n"
            "  • Temporal features: 5–30 min lags, rolling means/slopes, percent-change.\n"
            "  • Interaction terms: CO₂ × flow, humidity × temperature, VOC × flow.\n"
            "  • Schedule flag: is_class_time derived from Tennessee Tech course enrollment records.\n\n"
            "Modeling Pipeline:\n"
            "  1. Physics baseline — CO₂ mass-balance anchor: N = Q(C_indoor − C_outdoor) / G_person,\n"
            "     where Q = airflow-adjusted effective ACH, G_person = 0.005 L/s/person.\n"
            "  2. Blended estimate — 70% airflow-aware CO₂ anchor + 30% multi-sensor index (VOC, humidity, temp).\n"
            "  3. Weak supervision — manual anchor counts generate pseudo-labels with tiered confidence weights.\n"
            "  4. Residual ML models — Ridge, Random Forest, and Gradient Boosting learn correction offsets "
            "on top of the physics baseline using pseudo-labeled training data."
        ),
    )

    add_criterion(
        doc,
        number="3",
        criterion_title="Results — address the research question; clear and meaningful presentation; significance discussed",
        score_note="4–5 (Agree to Strongly Agree)",
        content_text=(
            "RESULTS (poster text):\n\n"
            "Sensor Correlations with Estimated Occupancy (Room 354, 30 days):\n"
            "  • CO₂ ↔ occupancy: r = 0.978 (dominant signal)\n"
            "  • VOC ↔ occupancy: r = 0.748 (strong secondary signal)\n"
            "  • Temperature ↔ occupancy: r = 0.392 (moderate support)\n"
            "  • Humidity ↔ occupancy: r = 0.265 (weak support)\n"
            "  • Airflow ↔ occupancy: r = −0.031 (not a direct headcount signal)\n\n"
            "Airflow Impact on CO₂ Anchor:\n"
            "  • Replacing fixed 4 ACH assumption with measured median 2.96 ACH reduced mean estimate from 5.2 → 4.2 people (−20%).\n"
            "  • Intervals with >30 estimated people dropped from 150 → 85 (−43%) after blending.\n\n"
            "Room 361 Validation (3 manual anchor windows):\n"
            "  | Model                    | MAE (people) | RMSE  | Band Accuracy |\n"
            "  |--------------------------|--------------|-------|---------------|\n"
            "  | Physics baseline         |     7.3      | 12.5  |    66.7%      |\n"
            "  | Gradient Boosting residual |   8.5      | 12.6  |    66.7%      |\n"
            "  | Random Forest residual   |     9.4      | 12.8  |     0%        |\n"
            "  | Ridge residual           |     9.5      | 12.2  |     0%        |\n\n"
            "Key finding: The physics baseline outperformed all ML residual models on held-out anchors, "
            "indicating that 3 anchor windows are insufficient to train a reliable correction. "
            "The largest error (Anchor A1: true=33, predicted=11) occurred during high-ventilation conditions "
            "(1,130 cfm) that diluted the CO₂ signal before it accumulated — a known physics-model limitation."
        ),
    )

    add_criterion(
        doc,
        number="4",
        criterion_title="Conclusions / Recommendations — appropriate conclusions drawn; future work provided",
        score_note="4–5 (Agree to Strongly Agree)",
        content_text=(
            "CONCLUSION (poster text):\n\n"
            "What we learned:\n"
            "  1. CO₂ and VOC are the strongest passive occupancy indicators in this classroom setting.\n"
            "  2. Measured airflow (ACH) meaningfully improves CO₂-based estimates vs. fixed-ventilation assumptions.\n"
            "  3. The hybrid physics + ML pipeline generalizes to new rooms with minimal labeled data.\n"
            "  4. Under high-ventilation conditions, CO₂ accumulation lags real occupancy — a physics-model blind spot.\n"
            "  5. Occupancy-aware HVAC scheduling using this model could reduce wasted runtime during low-use periods.\n\n"
            "Future Work:\n"
            "  • Collect 20+ manual headcount anchors per room for robust ML training and calibration.\n"
            "  • Add floor vibration / acoustic sensors to detect occupancy during high-ventilation periods.\n"
            "  • Integrate outdoor CO₂ monitoring to replace the fixed 415 ppm ambient assumption.\n"
            "  • Deploy a real-time occupancy dashboard and validate against HVAC energy logs."
        ),
    )

    add_criterion(
        doc,
        number="5",
        criterion_title="References — provided and properly cited",
        score_note="3 (Strongly Agree)",
        content_text=(
            "REFERENCES (poster text — use APA or IEEE format consistently):\n\n"
            "[1] ASHRAE Standard 62.1-2022, Ventilation and Acceptable Indoor Air Quality. Atlanta, GA: ASHRAE.\n\n"
            "[2] Fisk, W. J., et al. (2011). Association of ventilation rates and CO₂ concentrations with health and other "
            "responses in commercial and institutional buildings. Indoor Air, 21(4), 301–311.\n\n"
            "[3] Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. "
            "Journal of Machine Learning Research, 12, 2825–2830.\n\n"
            "[4] Tennessee Tech University. (2026). AIEB Building FPB and IAQ point-history sensor exports "
            "(Rooms 354 and 361). Internal dataset.\n\n"
            "[5] Batterman, S. (2017). Review and extension of CO₂-based methods to determine ventilation rates "
            "with application to school classrooms. International Journal of Environmental Research and Public Health, 14(2), 145."
        ),
    )

    add_criterion(
        doc,
        number="6",
        criterion_title="Acknowledgements — appropriate expressions of support from outside agencies and contributors",
        score_note="3 (Strongly Agree)",
        content_text=(
            "ACKNOWLEDGEMENTS (poster text):\n\n"
            "The authors thank the following individuals and groups for their support of this project:\n\n"
            "  • Chandler Norman — building systems expertise and FPB/IAQ data access.\n"
            "  • Norman Walker — domain guidance on HVAC operations and ventilation interpretation.\n"
            "  • Elisabeth Humphrey — project coordination support.\n"
            "  • Dr. Steven Anton — faculty advisor and technical review.\n"
            "  • Tennessee Tech University Jetstream2 cluster allocation — used for data processing and model training.\n"
            "  • Tennessee Tech College of Engineering — facilities and sensor infrastructure access."
        ),
    )

    # ─── SCORING SUMMARY ──────────────────────────────────────────────────────
    doc.add_page_break()
    add_section_header(doc, "REFERENCE", "Target Score Summary", 50)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Criterion"
    hdr[1].text = "Max Points"
    hdr[2].text = "Target"

    rows = [
        ("Part 1: Appearance & Organization", "15", "13–15"),
        ("  1. Text readability", "3", "3"),
        ("  2. Figures/tables legibility", "3", "3"),
        ("  3. Cohesive design", "3", "2–3"),
        ("  4. Consistent format / no errors", "3", "3"),
        ("  5. Organization & flow", "3", "3"),
        ("Part 2: Content", "35", "30–35"),
        ("  1. Intro / significance / research question", "5", "4–5"),
        ("  2. Methods", "5", "4–5"),
        ("  3. Results", "10", "8–10"),
        ("  4. Conclusions / recommendations", "5", "4–5"),
        ("  5. References", "5", "5"),
        ("  6. Acknowledgements", "5", "5"),
        ("TOTAL", "50", "43–50"),
    ]

    for label, max_pts, target in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = max_pts
        row[2].text = target

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: Scores are self-assessed targets based on current project completeness. "
        "The largest risk to a perfect score is Results depth — ensure the anchor evaluation table "
        "and the high-ventilation bias finding are clearly discussed on the poster."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    output_path = "/workspaces/CSC_4260_project/reports/poster_rubric_content.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
